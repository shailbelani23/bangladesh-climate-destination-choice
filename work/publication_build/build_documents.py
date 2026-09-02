#!/usr/bin/env python3
"""Build the final manuscript and one-page claim sheet as polished DOCX files."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "publication/manuscript/manuscript.md"
CLAIM = ROOT / "publication/claim_sheet.md"
OUT = ROOT / "publication/manuscript/bangladesh_climate_destination_choice_manuscript.docx"
CLAIM_OUT = ROOT / "publication/bangladesh_climate_destination_choice_claim_sheet.docx"

NAVY = RGBColor(23, 50, 79)
TEAL = RGBColor(11, 135, 147)
ORANGE = RGBColor(217, 107, 43)
MUTED = RGBColor(92, 112, 135)
LIGHT = "EAF1F5"
GRID = "CBD6E0"


def set_font(run, name="Aptos", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc, compact=False):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.5 if compact else 10.5)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(3 if compact else 6)
    normal.paragraph_format.line_spacing = 1.08 if compact else 1.22

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 7),
        ("Heading 2", 13, 12, 5),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.5 if compact else 10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(2 if compact else 4)
        style.paragraph_format.line_spacing = 1.08 if compact else 1.15

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Aptos"
    cap.font.size = Pt(8.5)
    cap.font.color.rgb = MUTED
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = False

    if "Table Caption" not in styles:
        tcap = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        tcap = styles["Table Caption"]
    tcap.font.name = "Aptos"
    tcap.font.size = Pt(9)
    tcap.font.bold = True
    tcap.font.color.rgb = NAVY
    tcap.paragraph_format.space_before = Pt(7)
    tcap.paragraph_format.space_after = Pt(4)
    tcap.paragraph_format.keep_with_next = True

    footer = section.footer
    add_page_number(footer.paragraphs[0])
    header = section.header.paragraphs[0]
    header.text = "CLIMATE-RELATED DESTINATION CHOICE IN BANGLADESH"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_font(run, size=8, bold=True, color=MUTED)


def add_inline(paragraph, text, size=None):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            if size:
                set_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Aptos Mono", size=size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if size:
            set_font(run, size=size)


def add_table(doc, rows):
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    if ncols == 7:
        widths = [2200, 1250, 850, 850, 850, 850, 2500]
    elif ncols == 6:
        widths = [2250, 1800, 950, 1050, 1050, 2260]
    else:
        widths = [9360 // ncols] * ncols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            mark_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_font(run, size=7.7 if ncols >= 7 else 8.3, bold=(r_idx == 0), color=NAVY)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 1
    spacer.add_run(" ")
    return table


def add_figure(doc, source, alt):
    path = (SRC.parent / source).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt.split(".")[0])
    cap = doc.add_paragraph(style="Figure Caption")
    add_inline(cap, alt)


def build_manuscript():
    lines = SRC.read_text().splitlines()
    doc = Document()
    configure_styles(doc)

    # Editorial-cover header pattern with restrained research styling.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("WHERE DO CLIMATE-AFFECTED HOUSEHOLDS GO?")
    set_font(r, name="Aptos Display", size=26, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Revealed District Choice in Bangladesh")
    set_font(r, name="Aptos Display", size=16, color=TEAL)
    for text, bold in [
        ("Shail Belani", True),
        ("Northwestern University, Evanston, Illinois, USA", False),
        ("shailbelani2027@u.northwestern.edu", False),
        ("September 3, 2026", False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        set_font(r, size=10.5, bold=bold, color=NAVY if bold else MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("ABSTRACT")
    set_font(r, size=10, bold=True, color=ORANGE)
    abstract_idx = lines.index("## Abstract")
    abstract_text = lines[abstract_idx + 2]
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(8)
    add_inline(p, abstract_text, size=9.2)
    p = doc.add_paragraph()
    add_inline(p, "Keywords: climate mobility; destination choice; Bangladesh; riverbank erosion; GIS; out-of-sample prediction", size=8.8)
    p.runs[0].bold = True
    doc.add_page_break()

    # Resume after abstract paragraph.
    lines = lines[abstract_idx + 3:]
    i = 0
    equation = False
    eq_lines = []
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "\\[":
            equation = True
            eq_lines = []
            i += 1
            continue
        if equation:
            if line == "\\]":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(" ".join(eq_lines).replace("\\", ""))
                set_font(r, name="Cambria Math", size=10.5, italic=True, color=NAVY)
                equation = False
            else:
                eq_lines.append(line)
            i += 1
            continue
        if line.startswith("# Appendix"):
            doc.add_page_break()
            p = doc.add_paragraph("Appendix", style="Heading 1")
            p.runs[0].font.size = Pt(20)
            i += 1
            continue
        if line.startswith("## "):
            heading = doc.add_heading(line[3:], level=1)
            if line == "## 9. Statements and declarations":
                heading.paragraph_format.page_break_before = True
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if re.match(r"^\*\*Table \d+\.", line):
            cap = doc.add_paragraph(style="Table Caption")
            add_inline(cap, line.strip("*"))
            table_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            parsed = []
            for idx, row in enumerate(table_lines):
                vals = [x.strip() for x in row.strip("|").split("|")]
                if idx == 1 and all(set(v) <= set("-:") for v in vals):
                    continue
                parsed.append(vals)
            add_table(doc, parsed)
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            add_figure(doc, m.group(2), m.group(1))
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            i += 1
            continue
        if line == "## References":
            doc.add_heading("References", level=1)
            i += 1
            continue
        p = doc.add_paragraph()
        if any(h.text == "References" for h in doc.paragraphs[-4:-1]) and not line.startswith("#"):
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(5)
        add_inline(p, line)
        i += 1

    # Keep references compact based on section position.
    in_refs = False
    for p in doc.paragraphs:
        if p.text == "References":
            in_refs = True
            continue
        if p.text == "Appendix":
            in_refs = False
        if in_refs and p.style.name == "Normal":
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(8)

    doc.core_properties.title = "Where Do Climate-Affected Households Go?"
    doc.core_properties.subject = "Revealed district choice in Bangladesh"
    doc.core_properties.author = "Shail Belani"
    doc.core_properties.keywords = "Bangladesh, climate mobility, destination choice, GIS"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


def build_claim_sheet():
    doc = Document()
    configure_styles(doc, compact=True)
    # Named one-page override: 9.5 pt body, 1.08 spacing, reduced heading spacing.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("FROZEN CLAIM SHEET")
    set_font(r, name="Aptos Display", size=21, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Where Do Climate-Affected Households Go?")
    set_font(r, size=11.5, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    add_inline(p, "Shail Belani | Undergraduate Researcher, Northwestern University", size=8.5)
    sections = [
        ("Main claim", "Across two independent Bangladesh surveys, four pre-specified district GIS measures improve out-of-sample destination probabilities relative to the same fitted gravity model. The evidence supports a predictive claim about where observed movers go. It does not identify a causal effect of any GIS characteristic on migration."),
        ("Different question", "Prior BEMP research estimates whether shocks change migration likelihood, type, and distance. This project conditions on an observed move and predicts the destination district. BIHS supplies the independent replication and main transport tests; BEMP supplies stronger shock timing."),
        ("Primary estimates", "BIHS national migrants: 1,857 events, gain 0.108 [0.082, 0.135]. BIHS erosion moves: 123 events, gain 0.098 [0.028, 0.163]. BEMP shock-linked moves: 184 events, gain 0.108 [0.023, 0.189]. Positive gain equals gravity log loss minus GIS log loss."),
        ("Generalization boundary", "The national BIHS model transfers to origins omitted from training: full-choice gain 0.101 [0.074, 0.128], interdistrict gain 0.107 [0.073, 0.144]. BIHS erosion moves do not support the same full-choice transfer: -0.016 [-0.086, 0.045]. Their interdistrict gain remains positive at 0.058 [-0.034, 0.140]. BEMP shows the same directional split."),
        ("Human illustration", "An anonymized household head reported moving from Faridpur to Manikganj in 2010 after erosion-related land loss. Gravity assigned Manikganj 7.0% probability and rank 6 of 64. GIS assigned 13.7% and rank 2. The case illustrates prediction; it does not reveal private reasoning."),
        ("Permitted language", "Use predicts, improves held-out probability, is associated with, and transfers across held-out origins. Do not use causes, proves preference, climate refugee, nationally representative climate displacement, or village-level claims. Public origins and destinations are districts."),
    ]
    for title, body in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title.upper())
        set_font(r, size=8.3, bold=True, color=ORANGE)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, body, size=9.3)
    doc.core_properties.title = "Frozen claim sheet: Bangladesh climate destination choice"
    doc.core_properties.author = "Shail Belani"
    CLAIM_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(CLAIM_OUT)


if __name__ == "__main__":
    build_manuscript()
    build_claim_sheet()
    print(OUT)
    print(CLAIM_OUT)
